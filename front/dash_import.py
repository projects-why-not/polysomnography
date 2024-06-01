import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import plotly.graph_objects as go
from docx.shared import Inches
import os
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
from dash import Dash, dcc, html, callback, Input, Output, clientside_callback

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["top", "left", "bottom", "right"]:
        border = kwargs.get(border_name)
        if border:
            element = OxmlElement(f"w:{border_name}")
            for attr, value in border.items():
                element.set(qn(f"w:{attr}"), str(value))
            tcPr.append(element)

def auto_shablon_generator(episodes_df, duration, processing_time):
    # Получение текущего времени и форматирование в строку
    dir_path = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")

    # Создание директории с именем текущего времени
    os.makedirs(dir_path)

    # Paths for output files
    output_word_path = dir_path + '/polysomnography_report.docx'
    output_pdf_path = dir_path + '/polysomnography_report.pdf'

    doc = Document()

    # Title
    title = doc.add_heading('Отчет для врача ... по полисомнографической записи ...', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add technical report section
    doc.add_heading('Технический отчет', level=2)
    doc.add_paragraph(f'Проведен анализ ПСГ длительностью {duration} минут')
    doc.add_paragraph(f'Оценивались сигналы типа ЭЭГ')
    doc.add_paragraph(f'Время обработки составило {processing_time} секунд')

    # Add clinical report section
    doc.add_heading('Клинический отчет', level=2)

    if episodes_df.empty:
        breathing_disorder_presence = ''
        doc.add_paragraph(f'По результатам анализа выявлено отсутствие признаков нарушений дыхания во сне')
    else:
        doc.add_paragraph(f'По результатам анализа выявлено наличие признаков нарушений дыхания во сне')
        doc.add_paragraph(f'За время ПСГ выявлено {len(episodes_df)} эпизодов нарушения дыхания (НРД)')

        # Add table
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ['№ эпизода НРД', 'Время начала регистрации эпизода НРД, с', 'Время завершения регистрации НРД, с',
                   'Длительность эпизода НРД, с', 'Тип эпизода НРД']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_border(hdr_cells[i], top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                            bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                            left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                            right={"sz": 12, "val": "single", "color": "000000", "space": "0"})

        for index, row in episodes_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["number"])
            row_cells[1].text = str(row["start_time"])
            row_cells[2].text = str(row["end_time"])
            row_cells[3].text = str(row["duration"])
            row_cells[4].text = row["type"]
            for cell in row_cells:
                set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                                bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                                left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                                right={"sz": 12, "val": "single", "color": "000000", "space": "0"})

        doc.add_heading('Характеристика длительности эпизодов нарушения дыхания', level=2)

        # Длительность эпизода НРД, с
        fig = go.Figure()
        fig.add_trace(go.Box(y=episodes_df['duration']))
        fig.update_xaxes(showticklabels=False)
        fig.update_layout(template='plotly', title={
            'text': "Длительность эпизода НРД, с",
            'y': 0.9,
            'x': 0.5,
            'xanchor': 'center',
            'yanchor': 'top',
        }, )
        fig.write_image(dir_path + '/image_graf.png')

        doc.add_picture(dir_path + '/image_graf.png', width=Inches(6))

        # Средняя длительность – […] секунд
        doc.add_paragraph(f'Средняя длительность – {round(episodes_df['duration'].mean(), 2)} секунд')
        # Медиана длительности НРД -- […] секунд
        doc.add_paragraph(f'Медиана длительности НРД – {round(episodes_df['duration'].median(), 2)} секунд')
        # Разброс длительности НРД -- от […] секунд до […] секунд
        doc.add_paragraph(
            f'Разброс длительности НРД - от {round(episodes_df['duration'].quantile(0.25), 2)} секунд до {round(episodes_df['duration'].quantile(0.75), 2)} секунд')
        # round(float(mean_duration), 2), round(float(median_duration), 2), spread_duration

        doc.add_heading('Анализ эпизодов нарушения дыхания в зависимости от времени регистрации', level=2)

        # Количество разных типов апноэ за всю ночь
        fig = go.Figure()
        fig.add_trace(go.Bar(x=episodes_df[episodes_df["type"].str.contains("апноэ")]['type'].unique(),
                             y=episodes_df[episodes_df["type"].str.contains("апноэ")]['type'].value_counts()))
        fig.update_layout(template='plotly', title={
            'text': "Количество разных типов апноэ за всю ночь",
            'y': 0.9,
            'x': 0.5,
            'xanchor': 'center',
            'yanchor': 'top',
        }, )
        fig.write_image(dir_path + '/image_graf1.png')

        doc.add_picture(dir_path + '/image_graf1.png', width=Inches(6))

        values_of_different_types_of_respiratory_disorders = episodes_df["type"].replace({
            "центральное апноэ": "апноэ",
            "обструктивное апноэ": "апноэ",
            "апноэ": "апноэ",
            "апноэ неопределенного типа": "апноэ"
        })

        # Количество разных типов нарушений дыхания
        fig = go.Figure()
        fig.add_trace(go.Bar(x=values_of_different_types_of_respiratory_disorders.unique(),
                             y=values_of_different_types_of_respiratory_disorders.value_counts()))
        fig.update_layout(template='plotly', title={
            'text': "Количество разных типов нарушений дыхания",
            'y': 0.9,
            'x': 0.5,
            'xanchor': 'center',
            'yanchor': 'top',
        }, )
        fig.write_image(dir_path + '/image_graf2.png', scale=2)

        doc.add_picture(dir_path + '/image_graf2.png', width=Inches(6))

        # Распределение эпизодов НРД по первой, второй и последней третях ночного сна
        fig = go.Figure()
        fig.add_trace(go.Bar(x=['первая треть', 'вторая треть', 'последняя треть'],
                             y=[len(episodes_df[episodes_df['start_time'] < duration / 3]),
                                len(episodes_df[(episodes_df['start_time'] > duration / 3) & (
                                            episodes_df['start_time'] < duration * 2 / 3)]),
                                len(episodes_df[episodes_df['start_time'] > duration * 2 / 3])]))
        fig.update_layout(template='plotly', title={
            'text': "Распределение эпизодов НРД по третям ночного сна",
            'y': 0.9,
            'x': 0.5,
            'xanchor': 'center',
            'yanchor': 'top',
        }, )
        fig.write_image(dir_path + '/image_graf3.png', scale=2)

        doc.add_picture(dir_path + '/image_graf3.png', width=Inches(6))

    # Save the document
    doc.save(output_word_path)

    # Convert the document to PDF
    convert(output_word_path, output_pdf_path)

    return output_word_path, output_pdf_path

def div_maker(fig):
    return html.Div(
        children=dcc.Graph(figure=fig),
        style={"display": "flex", "justify-content": "center"},
    )

def upper_plot_generator(data_poly, data_gipno):

    i = 0

    fig = go.Figure()
    for column in data_poly.columns[:-1]:
        if column in ['Airflow', 'Chest', 'Abdomen']:
            fig.add_trace(go.Scatter(x=data_poly['Time'], y=data_poly[column] - 0.005 * i, name=column))
            i += 1

    if data_gipno is not None:
        fig.add_trace(go.Scatter(x=data_gipno['Time'], y=data_gipno['stage'] / 900 - 0.0075 * i, name='gipno'))

    fig.update_yaxes(showticklabels=False)
    fig.update_xaxes(
        tickformat="%H:%M:%S"
    )
    fig.update_xaxes(rangeslider_visible=True)  # ranges slider

    return div_maker(fig)

def lower_plot_generator(data):
    i = 0

    fig = go.Figure()
    for column in data.columns[:-1]:
        if column in ['Fp1-M2', 'C3-M2', 'O1-M2', 'Fp2-M1', 'C4-M1', 'O2-M1']:
            fig.add_trace(go.Scatter(x=data['Time'], y=data[column] - 0.0005 * i, name=column))
            i += 1
    fig.update_yaxes(showticklabels=False)
    fig.update_xaxes(
        tickformat="%H:%M:%S"
    )
    fig.update_xaxes(rangeslider_visible=True)  # ranges slider

    return div_maker(fig)

def start_dash(data_poly, data_gipno, episodes, hertz, processing_time):
    data_poly, data_gipno, duration = data_preprocessing(data_poly, data_gipno, hertz)

    path_word, path_pdf = auto_shablon_generator(episodes, duration, processing_time)

    app = Dash(__name__)
    app._favicon = "images/icon3.png"
    app.title = "Результат анализа"

    title_page = html.H1(children='Результаты анализа сна по ЭЭГ', className='title')

    link_1 = html.A(children=html.Div(children=html.P('Загрузить отчёт .docx', className='download_button_lable'), className='link__button'), href='link/to/your/download/file', download=path_word)
    link_2 = html.A(children=html.Div(children=html.P('Загрузить отчёт .pdf', className='download_button_lable'), className='link__button'), href='link/to/your/download/file', download=path_pdf)
    link_place = html.Div(children=[link_1, link_2], className='link__place')

    report_message = html.H3('Загрузить полный отчёт на компьютер', className='report_message')
    report_place = html.Div(children=[report_message, link_place], className='report_place')

    app.layout = html.Div(children=[
        title_page,
        html.Div(children=upper_plot_generator(data_poly, data_gipno), className='upper_plot'),
        html.Div(children='Показать ЭЭГ',className="lower-plot__button", id="test"),
        html.Div(children=lower_plot_generator(data_poly), className='lower_plot', id='lowerplot'),
        report_place],
        className='page')
    
    clientside_callback(
    """
    function(id) {
        const el = document.getElementById("test");
        const tosee = document.getElementById("lowerplot");
        el.addEventListener('click', () => {
        tosee.classList.toggle("lower_plot_visible")
        if (el.textContent === "Показать ЭЭГ"){
        el.textContent = "Скрыть ЭЭГ"
        el.style.marginBottom ='0'
        } else {
        el.textContent = "Показать ЭЭГ"
        el.style.marginBottom = '32px'
        }
        })
        
        
        
        return window.dash_clientside.no_update
    }
    """,
    Output("test","id"),
    Input("test","id"),
    )

    app.run_server(debug=True)

def data_preprocessing(data_poly, data_gipno, hertz):
    poly = pd.DataFrame()
    for column in data_poly.columns:
        poly[column] = data_poly[column].rolling(window=hertz).mean()[::hertz]
    poly.loc[0, poly.columns] = data_poly.loc[0, poly.columns]
    poly['Time'] = pd.date_range(start='00:00:00', periods=len(poly), freq='1s')


    data_gipno['Time'] = pd.date_range(start='00:00:00', periods=len(data_gipno), freq='30s')
    data_gipno = data_gipno[['Time', 'stage']].copy()

    duration = int((poly['Time'].iloc[-1] - poly['Time'].iloc[0]).total_seconds())

    return poly, data_gipno, duration

if __name__ == "__main__":
    #input
    data1 = pd.read_csv('full_data.csv')
    data2 = pd.read_csv('hypno.csv')
    hertz = 200
    processing_time = random.randrange(5,15)
    episodes_data = pd.read_csv('episodes.csv')
    #func
    start_dash(data1, data2, episodes_data, hertz, processing_time)

